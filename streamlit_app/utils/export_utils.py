# ================================================================
#  Employee Management System (EMS)
#  Developed by: Sam Ranjith Paul
#  GitHub: https://github.com/samranjithpaul
#  LinkedIn: https://www.linkedin.com/in/Samranjithpaul
#  Unauthorized removal of this header is prohibited.
# ================================================================
"""
Export Utilities for Employee Management System
Handles Excel, Word, and PDF exports for employee data
"""
import os
import pandas as pd
from datetime import datetime
from typing import Dict, List, Any, Optional
from docx import Document
from docx.shared import Inches, Pt
from docx.enum.text import WD_ALIGN_PARAGRAPH
from fpdf import FPDF
import zipfile
import tempfile

from .api_client import BASE_URL, get_headers, get_employees, get_employee, get_employment_history


# Create exports directory if it doesn't exist
EXPORTS_DIR = os.path.join(os.path.dirname(os.path.dirname(__file__)), "exports")
os.makedirs(EXPORTS_DIR, exist_ok=True)


def get_timestamp() -> str:
    """Get current timestamp in YYYYMMDD format"""
    return datetime.now().strftime("%Y%m%d")


def export_all_to_excel(token: str) -> Optional[str]:
    """
    Export all employees to Excel format
    
    Returns:
        Path to generated Excel file or None if error
    """
    try:
        # Get all employees from API
        employees = get_employees(token)
        
        if not employees:
            return None
        
        # Prepare data for DataFrame
        excel_data = []
        for emp in employees:
            excel_data.append({
                "Employee ID": emp.get("emp_id", ""),
                "First Name": emp.get("first_name", ""),
                "Last Name": emp.get("last_name", ""),
                "Full Name": f"{emp.get('first_name', '')} {emp.get('last_name', '')}".strip(),
                "Email": emp.get("email") or "",
                "Phone": emp.get("phone") or "",
                "Department": emp.get("department") or "",
                "Designation": emp.get("designation") or "",
                "Joining Date": emp.get("joining_date") or "",
                "Status": emp.get("status", ""),
                "Created At": emp.get("created_at", ""),
            })
        
        # Create DataFrame
        df = pd.DataFrame(excel_data)
        
        # Generate filename
        filename = f"employee_data_{get_timestamp()}.xlsx"
        filepath = os.path.join(EXPORTS_DIR, filename)
        
        # Write to Excel with formatting
        with pd.ExcelWriter(filepath, engine='openpyxl') as writer:
            df.to_excel(writer, sheet_name='Employees', index=False)
            
            # Get the workbook and worksheet
            workbook = writer.book
            worksheet = writer.sheets['Employees']
            
            # Auto-adjust column widths
            for column in worksheet.columns:
                max_length = 0
                column_letter = column[0].column_letter
                for cell in column:
                    try:
                        if len(str(cell.value)) > max_length:
                            max_length = len(str(cell.value))
                    except:
                        pass
                adjusted_width = min(max_length + 2, 50)
                worksheet.column_dimensions[column_letter].width = adjusted_width
            
            # Format header row
            from openpyxl.styles import Font, PatternFill, Alignment
            
            header_fill = PatternFill(start_color="366092", end_color="366092", fill_type="solid")
            header_font = Font(bold=True, color="FFFFFF", size=11)
            
            for cell in worksheet[1]:
                cell.fill = header_fill
                cell.font = header_font
                cell.alignment = Alignment(horizontal="center", vertical="center")
        
        return filepath
        
    except Exception as e:
        print(f"Error exporting to Excel: {e}")
        return None


def export_employee_to_word(emp_id: int, token: str) -> Optional[str]:
    """
    Export single employee to Word document
    
    Args:
        emp_id: Employee ID
        token: JWT authentication token
        
    Returns:
        Path to generated Word file or None if error
    """
    try:
        # Get employee details
        emp = get_employee(token, emp_id)
        if not emp:
            return None
        
        # Get employment history
        history = get_employment_history(token, emp_id)
        
        # Create Word document
        doc = Document()
        
        # Set document margins
        sections = doc.sections
        for section in sections:
            section.top_margin = Inches(0.5)
            section.bottom_margin = Inches(0.5)
            section.left_margin = Inches(1)
            section.right_margin = Inches(1)
        
        # Title
        title = doc.add_heading(f"Employee Profile", 0)
        title.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Employee Name
        name_para = doc.add_paragraph()
        name_run = name_para.add_run(f"{emp.get('first_name', '')} {emp.get('last_name', '')}".strip())
        name_run.font.size = Pt(20)
        name_run.bold = True
        name_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        doc.add_paragraph()  # Spacing
        
        # Basic Information Section
        doc.add_heading("Basic Information", level=1)
        
        info_items = [
            ("Employee ID", str(emp.get("emp_id", ""))),
            ("Department", emp.get("department") or "N/A"),
            ("Designation", emp.get("designation") or "N/A"),
            ("Email", emp.get("email") or "N/A"),
            ("Phone", emp.get("phone") or "N/A"),
            ("Joining Date", str(emp.get("joining_date") or "N/A")),
            ("Status", emp.get("status", "N/A")),
        ]
        
        for label, value in info_items:
            para = doc.add_paragraph()
            para.add_run(f"{label}: ").bold = True
            para.add_run(value)
        
        doc.add_paragraph()  # Spacing
        
        # Employment History Section
        if history:
            doc.add_heading("Employment History", level=1)
            for idx, hist in enumerate(history, 1):
                doc.add_paragraph(f"{idx}. {hist.get('company_name', 'N/A')}", style='List Bullet')
                para = doc.add_paragraph()
                para.add_run("   Position: ").bold = True
                para.add_run(hist.get("position", "N/A"))
                para = doc.add_paragraph()
                para.add_run("   Duration: ").bold = True
                para.add_run(f"{hist.get('start_date', 'N/A')} to {hist.get('end_date', 'N/A')}")
                doc.add_paragraph()  # Spacing
        else:
            doc.add_heading("Employment History", level=1)
            doc.add_paragraph("No previous employment history recorded.")
        
        # Footer with developer credit
        doc.add_paragraph()  # Spacing
        footer_para = doc.add_paragraph()
        footer_text = f"Generated by Employee Management System | Developed by Sam Ranjith Paul | github.com/samranjithpaul | linkedin.com/in/Samranjithpaul | {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}"
        footer_para.add_run(footer_text).italic = True
        footer_para.alignment = WD_ALIGN_PARAGRAPH.CENTER
        
        # Generate filename
        first_name = emp.get("first_name", "Employee").replace(" ", "_")
        last_name = emp.get("last_name", "").replace(" ", "_")
        filename = f"Employee_{first_name}_{last_name}_Profile.docx"
        filepath = os.path.join(EXPORTS_DIR, filename)
        
        # Save document
        doc.save(filepath)
        
        return filepath
        
    except Exception as e:
        print(f"Error exporting to Word: {e}")
        return None


def export_employee_to_pdf(emp_id: int, token: str) -> Optional[str]:
    """
    Export single employee to PDF format
    
    Args:
        emp_id: Employee ID
        token: JWT authentication token
        
    Returns:
        Path to generated PDF file or None if error
    """
    try:
        # Get employee details
        emp = get_employee(token, emp_id)
        if not emp:
            return None
        
        # Get employment history
        history = get_employment_history(token, emp_id)
        
        # Create PDF
        pdf = FPDF()
        pdf.add_page()
        
        # Set font
        pdf.set_font("Arial", "B", 20)
        
        # Title
        pdf.cell(0, 15, "Employee Profile", ln=True, align="C")
        pdf.ln(5)
        
        # Employee Name
        full_name = f"{emp.get('first_name', '')} {emp.get('last_name', '')}".strip()
        pdf.set_font("Arial", "B", 16)
        pdf.cell(0, 10, full_name, ln=True, align="C")
        pdf.ln(10)
        
        # Draw line
        pdf.line(10, pdf.get_y(), 200, pdf.get_y())
        pdf.ln(5)
        
        # Basic Information
        pdf.set_font("Arial", "B", 14)
        pdf.cell(0, 10, "Basic Information", ln=True)
        pdf.ln(3)
        
        pdf.set_font("Arial", "", 11)
        
        info_items = [
            ("Employee ID", str(emp.get("emp_id", ""))),
            ("Department", emp.get("department") or "N/A"),
            ("Designation", emp.get("designation") or "N/A"),
            ("Email", emp.get("email") or "N/A"),
            ("Phone", emp.get("phone") or "N/A"),
            ("Joining Date", str(emp.get("joining_date") or "N/A")),
            ("Status", emp.get("status", "N/A")),
        ]
        
        for label, value in info_items:
            pdf.set_font("Arial", "B", 11)
            pdf.cell(60, 8, f"{label}:", ln=0)
            pdf.set_font("Arial", "", 11)
            pdf.cell(0, 8, value, ln=True)
        
        pdf.ln(5)
        pdf.line(10, pdf.get_y(), 200, pdf.get_y())
        pdf.ln(5)
        
        # Employment History
        pdf.set_font("Arial", "B", 14)
        pdf.cell(0, 10, "Employment History", ln=True)
        pdf.ln(3)
        
        if history:
            pdf.set_font("Arial", "", 11)
            for idx, hist in enumerate(history, 1):
                pdf.set_font("Arial", "B", 11)
                pdf.cell(0, 8, f"{idx}. {hist.get('company_name', 'N/A')}", ln=True)
                pdf.set_font("Arial", "", 10)
                pdf.cell(20, 6, "", ln=0)  # Indent
                pdf.cell(0, 6, f"Position: {hist.get('position', 'N/A')}", ln=True)
                pdf.cell(20, 6, "", ln=0)  # Indent
                pdf.cell(0, 6, f"Duration: {hist.get('start_date', 'N/A')} to {hist.get('end_date', 'N/A')}", ln=True)
                pdf.ln(3)
        else:
            pdf.set_font("Arial", "", 11)
            pdf.cell(0, 8, "No previous employment history recorded.", ln=True)
        
        # Footer
        pdf.set_y(-20)
        pdf.set_font("Arial", "I", 9)
        pdf.cell(0, 10, f"Generated by EMS System - {datetime.now().strftime('%Y-%m-%d %H:%M:%S')}", align="C")
        
        # Generate filename
        first_name = emp.get("first_name", "Employee").replace(" ", "_")
        last_name = emp.get("last_name", "").replace(" ", "_")
        filename = f"Employee_{first_name}_{last_name}_Profile.pdf"
        filepath = os.path.join(EXPORTS_DIR, filename)
        
        # Save PDF
        pdf.output(filepath)
        
        return filepath
        
    except Exception as e:
        print(f"Error exporting to PDF: {e}")
        return None


def export_all_to_pdf(token: str) -> Optional[str]:
    """
    Export all employees to a single multi-page PDF
    
    Returns:
        Path to generated PDF file or None if error
    """
    try:
        # Get all employees
        employees = get_employees(token)
        
        if not employees:
            return None
        
        # Create PDF
        pdf = FPDF()
        
        # Add a page for each employee
        for idx, emp in enumerate(employees, 1):
            # Add new page (except for first employee)
            if idx > 1:
                pdf.add_page()
            
            # Title
            pdf.set_font("Arial", "B", 18)
            pdf.cell(0, 15, "All Employees Report", ln=True, align="C")
            pdf.ln(5)
            
            # Employee Name
            full_name = f"{emp.get('first_name', '')} {emp.get('last_name', '')}".strip()
            pdf.set_font("Arial", "B", 16)
            pdf.cell(0, 10, f"Employee #{idx}: {full_name}", ln=True, align="C")
            pdf.ln(10)
            
            # Draw line
            pdf.line(10, pdf.get_y(), 200, pdf.get_y())
            pdf.ln(5)
            
            # Basic Information
            pdf.set_font("Arial", "B", 12)
            pdf.cell(0, 8, "Basic Information", ln=True)
            pdf.ln(3)
            
            pdf.set_font("Arial", "", 10)
            
            info_items = [
                ("Employee ID", str(emp.get("emp_id", ""))),
                ("Department", emp.get("department") or "N/A"),
                ("Designation", emp.get("designation") or "N/A"),
                ("Email", emp.get("email") or "N/A"),
                ("Phone", emp.get("phone") or "N/A"),
                ("Joining Date", str(emp.get("joining_date") or "N/A")),
                ("Status", emp.get("status", "N/A")),
            ]
            
            for label, value in info_items:
                pdf.set_font("Arial", "B", 10)
                pdf.cell(50, 7, f"{label}:", ln=0)
                pdf.set_font("Arial", "", 10)
                pdf.cell(0, 7, value, ln=True)
            
            pdf.ln(5)
            
            # Check if we need a new page for next employee
            if pdf.get_y() > 250 and idx < len(employees):
                continue  # Will add new page in next iteration
        
        # Footer on last page with developer credit
        pdf.set_y(-20)
        pdf.set_font("Arial", "I", 8)
        pdf.cell(0, 10, f"Generated by EMS | Developed by Sam Ranjith Paul | github.com/samranjithpaul | {datetime.now().strftime('%Y-%m-%d %H:%M:%S')} | Total Employees: {len(employees)}", align="C")
        
        # Generate filename
        filename = f"All_Employees_Report_{get_timestamp()}.pdf"
        filepath = os.path.join(EXPORTS_DIR, filename)
        
        # Save PDF
        pdf.output(filepath)
        
        return filepath
        
    except Exception as e:
        print(f"Error exporting all to PDF: {e}")
        return None


def export_all_pdfs_to_zip(token: str) -> Optional[str]:
    """
    Export all employees as individual PDFs and bundle them in a ZIP file
    
    Returns:
        Path to generated ZIP file or None if error
    """
    try:
        # Get all employees
        employees = get_employees(token)
        
        if not employees:
            return None
        
        # Create temporary directory for PDFs
        temp_dir = tempfile.mkdtemp()
        pdf_files = []
        
        # Generate PDF for each employee
        for emp in employees:
            emp_id = emp.get("emp_id")
            if emp_id:
                pdf_path = export_employee_to_pdf(emp_id, token)
                if pdf_path:
                    # Copy to temp directory
                    import shutil
                    temp_pdf_path = os.path.join(temp_dir, os.path.basename(pdf_path))
                    shutil.copy2(pdf_path, temp_pdf_path)
                    pdf_files.append(temp_pdf_path)
        
        if not pdf_files:
            return None
        
        # Create ZIP file
        zip_filename = f"Employee_PDFs_Archive_{get_timestamp()}.zip"
        zip_filepath = os.path.join(EXPORTS_DIR, zip_filename)
        
        with zipfile.ZipFile(zip_filepath, 'w', zipfile.ZIP_DEFLATED) as zipf:
            for pdf_file in pdf_files:
                zipf.write(pdf_file, os.path.basename(pdf_file))
        
        # Clean up temp directory
        import shutil
        shutil.rmtree(temp_dir, ignore_errors=True)
        
        return zip_filepath
        
    except Exception as e:
        print(f"Error creating ZIP archive: {e}")
        return None

